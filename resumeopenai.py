#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
import streamlit as st
import pdfplumber
import docx
import re
import openai
import json
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

if not openai.api_key:
    raise ValueError("OpenAI API key not found. Please set it in your .env file or system environment.")

# PDF Text Extraction
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    return text.strip()

# DOCX Text Extraction
def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs]).strip()

# Parse Resume Based on File Type
def parse_resume(file):
    if file.name.endswith(".pdf"):
        return extract_text_from_pdf(file), "pdf"
    elif file.name.endswith(".docx"):
        return extract_text_from_docx(file), "docx"
    else:
        return None, None

# Extract Sections from Resume
def extract_sections(resume_text):
    sections = {}
    current_section = "Other"
    lines = resume_text.split("\n")
    for line in lines:
        if re.match(r"^\s*[A-Z][A-Za-z ]+:\s*$", line):
            current_section = line.strip()
            sections[current_section] = []
        else:
            sections.setdefault(current_section, []).append(line)
    return sections

# Extract Skills from Job Description using OpenAI
def extract_skills_with_gpt(job_desc):
    prompt = f"""
    Extract key skills, tools, technologies, and relevant qualifications from the following job description.
    Classify them into resume sections such as 'Technical Skills', 'Certifications', 'Experience', 'Projects', etc.
    Provide the output as a JSON object where keys are section names and values are lists of skills.
    Job Description:
    {job_desc}
    """
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an AI expert in resume optimization."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2
        )
        extracted_text = response.choices[0].message.content.strip()
        extracted_data = json.loads(extracted_text)
        return extracted_data
    except json.JSONDecodeError:
        st.error("Error: Failed to parse extracted skills from OpenAI response.")
        return {}
    except openai.OpenAIError as e:
        st.error(f"Error extracting skills: {e}")
        return {}

# Get User Approval for Skills
def get_user_approval(extracted_data):
    approved_data = {}
    for section, items in extracted_data.items():
        approved_items = []
        for item in items:
            if st.checkbox(f"Add to {section}: {item}?"):
                approved_items.append(item)
        if approved_items:
            approved_data[section] = approved_items
    return approved_data

# Update Resume While Retaining Format
def update_resume(sections, approved_data):
    for section, items in approved_data.items():
        if section in sections:
            sections[section].append(", ".join(items))
        else:
            sections[section] = [", ".join(items)]
    return sections

# Convert Updated Sections to DOCX
def save_as_docx(sections, original_docx):
    doc = Document(original_docx)
    for para in doc.paragraphs:
        for section, content in sections.items():
            if para.text.strip() == section:
                para.add_run("\n" + "\n".join(content))
    output_path = "Updated_Resume.docx"
    doc.save(output_path)
    return output_path

# Convert Updated Sections to PDF
def save_as_pdf(sections):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for section, content in sections.items():
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(200, 10, txt=section, ln=True, align='L')
        pdf.set_font("Arial", size=12)
        for line in content:
            pdf.multi_cell(0, 10, txt=line)
    output_path = "Updated_Resume.pdf"
    pdf.output(output_path)
    return output_path

# Calculate ATS Score
def calculate_ats_score(text1, text2):
    vectorizer = CountVectorizer().fit_transform([text1, text2])
    vectors = vectorizer.toarray()
    return cosine_similarity([vectors[0]], [vectors[1]])[0][0] * 100

# Streamlit UI
st.title("🔍 AI-Powered Resume Enhancer with ATS Scoring")

uploaded_file = st.file_uploader("📂 Upload your resume", type=["pdf", "docx"])
job_description = st.text_area("📝 Paste the job description here")

if uploaded_file and job_description:
    resume_text, file_type = parse_resume(uploaded_file)
    
    if resume_text:
        sections = extract_sections(resume_text)
        extracted_data = extract_skills_with_gpt(job_description)
        
        if extracted_data:
            st.subheader("✅ Select skills to add:")
            approved_data = get_user_approval(extracted_data)
            
            updated_sections = update_resume(sections, approved_data)
            ats_score_old = calculate_ats_score(job_description, resume_text)
            updated_resume_text = "\n".join([f"{sec}\n" + "\n".join(content) for sec, content in updated_sections.items()])
            ats_score_new = calculate_ats_score(job_description, updated_resume_text)
            
            st.write(f"### 📊 ATS Score (Old Resume): {ats_score_old:.2f}%")
            st.write(f"### 📊 ATS Score (New Resume): {ats_score_new:.2f}%")
            
            if file_type == "docx":
                output_file = save_as_docx(updated_sections, uploaded_file)
            else:
                output_file = save_as_pdf(updated_sections)
            
            with open(output_file, "rb") as f:
                st.download_button("📥 Download Updated Resume", f, file_name=output_file)
    else:
        st.error("❌ Unable to extract text from the resume. Please check the file format.")


# In[ ]:




